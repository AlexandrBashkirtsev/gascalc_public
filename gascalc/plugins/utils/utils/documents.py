import os
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from gascalc import app
from .load import format_route

def get_docx(summary, settings, gen_time, mass_export, **kwargs):

	# variables
	generation_date = gen_time.date().strftime("%d.%m.%y")
	date = f'c {str(generation_date)} по {str(generation_date)}.'
	organization = settings['org_name']
	uadress = settings['org_address']
	OGRN = settings['OGRN']
	INN = settings['INN']
	car_info = settings['car']
	fullname = settings['driver_fullname']
	fuel = settings['fuel']
	distance = summary['summary']['distance']
	consumption = settings['consumption']
	fuel_marge = summary['summary']['fuel_marge']
	accountant = settings['accountant']

	formatted_route = format_route(summary=summary,
									start_time=gen_time,
									web=False)

	# constants
	title 	= 'Путевой лист легкового автомобиля N  ____'
	med_before = 'Предрейсовый медосмотр ______________  ________  ________________ '
	med_line = '_______/________________________/'
	med_after = 'Послерейсовый медосмотр _____________  ________  ________________ '

	odometr = 'Показания одометра при выезде - заезде на парковку'
	odo_line = '__________________ - _________________'

	#------------------STYLES-----------------

	if not mass_export:

		document = Document()

		styles = document.styles

		TP_STYLE = styles.add_style('TP_STYLE', WD_STYLE_TYPE.PARAGRAPH)
		FONT_TP_STYLE = TP_STYLE.font
		FONT_TP_STYLE.name = 'Times New Roman'
		FONT_TP_STYLE.size = Pt(12)

		HP_STYLE = styles.add_style('HP_STYLE', WD_STYLE_TYPE.PARAGRAPH)
		FONT_HP_STYLE = HP_STYLE.font
		FONT_HP_STYLE.name = 'Times New Roman'
		FONT_HP_STYLE.size = Pt(14)

		TABLE_STYLE = styles.add_style('TABLE_STYLE', WD_STYLE_TYPE.TABLE)
		FONT_TABLE_STYLE = TABLE_STYLE.font
		FONT_TABLE_STYLE.name = 'Times New Roman'
		FONT_TABLE_STYLE.size = Pt(10)	

	else:

		document = kwargs['document']
		TP_STYLE = kwargs['TP_STYLE']
		HP_STYLE = kwargs['HP_STYLE']
		TABLE_STYLE = kwargs['TABLE_STYLE']


	

	#------------------LAYOUT------------------
	document.add_paragraph(style=HP_STYLE)\
						.add_run(title)

	document.add_paragraph(style=HP_STYLE)\
						.add_run(date)

	document.add_paragraph(style=TP_STYLE)\
						.add_run('Организация: ' + organization)

	document.add_paragraph(style=TP_STYLE)\
						.add_run(uadress)

	document.add_paragraph(style=TP_STYLE)\
						.add_run('ОГРН ' + OGRN + ' ИНН ' + INN)

	document.add_paragraph(style=TP_STYLE)\
						.add_run('Легковой Автомобиль: ' + car_info)

	document.add_paragraph(style=TP_STYLE)\
						.add_run('ФИО водителя: ' + fullname)

	document.add_paragraph(style=TP_STYLE)\
						.add_run(med_before)

	document.add_paragraph(style=TP_STYLE)\
						.add_run(med_line)

	document.add_paragraph(style=TP_STYLE)\
						.add_run(med_after)

	document.add_paragraph(style=TP_STYLE)\
						.add_run(med_line)

	document.add_paragraph(style=TP_STYLE)\
						.add_run(odometr)

	document.add_paragraph(style=TP_STYLE)\
						.add_run(odo_line)

	document.add_paragraph(style=TP_STYLE)\
						.add_run(med_line).add_break()


	# info table
	table_info = document.add_table(rows=2, cols=5, style='Table Grid')
	table_info.style.font.name = 'Times New Roman'
	table_info.style.font.size = Pt(10)

	hdr_cells = table_info.rows[0].cells
	val_cells = table_info.rows[1].cells

	hdr_cells[0].text  = 'Наименование топлива'
	hdr_cells[1].text  = 'Пробег (км)'
	hdr_cells[2].text  = 'Норма расхода (л/100 км)'
	hdr_cells[3].text  = 'Расход топлива по норме (л)'
	hdr_cells[4].text  = 'Расход топлива по факту (л)'

	val_cells[0].text  = str(fuel)
	val_cells[1].text  = str(distance)
	val_cells[2].text  = str(consumption)
	val_cells[3].text  = "{:.2f}".format(fuel_marge)
	val_cells[4].text  = "{:.2f}".format(fuel_marge)


	for row in table_info.rows:
		for cell in row.cells:
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	document.add_paragraph()

	# route table
	table_route = document.add_table(rows=1, cols=5, style='Table Grid')
	table_route.allow_autofit = False
	table_route.autofit = False
	table_route.style.font.name = 'Times New Roman'
	table_route.style.font.size = Pt(10)

	hdr_cells = table_route.rows[0].cells

	hdr_cells[0].text  = '№ п/п'
	hdr_cells[1].text  = 'Маршрут'
	hdr_cells[2].text  = 'Время выезда (час:минута)'
	hdr_cells[3].text  = 'Время приезда (час:минута)'
	hdr_cells[4].text  = 'Пройдено, км'

	for location in formatted_route:
		row = table_route.add_row()
		row.cells[0].text = location['index']
		row.cells[1].text = location['location']
		row.cells[2].text = location['out_time'][:-3]
		row.cells[3].text = location['in_time'][:-3]
		row.cells[4].text = location['distance']

	for row in table_route.rows:
		for cell in row.cells:
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	table_route.columns[0].width = Inches(0.3)
	table_route.columns[1].width = Inches(2.8)
	table_route.columns[4].width = Inches(0.5)

	document.add_paragraph()

	document.add_paragraph(style=TP_STYLE)\
						.add_run('Итоговый пробег: ' + str(distance) + ' км')

	document.add_paragraph(style=TP_STYLE)\
						.add_run('Водитель		 ___________________ ('+ fullname +' )')

	document.add_paragraph(style=TP_STYLE)\
						.add_run('Бухгалтер		 ___________________ ('+ accountant +' )')

	document.add_page_break()

	if not mass_export:

		name = fullname + '_' + generation_date + '.docx'
		filename = os.path.join(app.config['UPLOAD_FOLDER'], name)
		document.save(filename)

		return name

	else:

		return document

def mass_gen(routes):

	document = Document()

	styles = document.styles

	TP_STYLE = styles.add_style('TP_STYLE', WD_STYLE_TYPE.PARAGRAPH)
	FONT_TP_STYLE = TP_STYLE.font
	FONT_TP_STYLE.name = 'Times New Roman'
	FONT_TP_STYLE.size = Pt(12)
	PARAGRAPH_TP_FORMAT = TP_STYLE.paragraph_format
	PARAGRAPH_TP_FORMAT.space_before = Pt(1)
	PARAGRAPH_TP_FORMAT.line_spacing = Pt(10)

	HP_STYLE = styles.add_style('HP_STYLE', WD_STYLE_TYPE.PARAGRAPH)
	FONT_HP_STYLE = HP_STYLE.font
	FONT_HP_STYLE.name = 'Times New Roman'
	FONT_HP_STYLE.size = Pt(14)
	PARAGRAPH_HP_FORMAT = HP_STYLE.paragraph_format
	PARAGRAPH_HP_FORMAT.space_before = Pt(1)
	PARAGRAPH_HP_FORMAT.line_spacing = Pt(12)

	TABLE_STYLE = styles.add_style('TABLE_STYLE', WD_STYLE_TYPE.TABLE)
	FONT_TABLE_STYLE = TABLE_STYLE.font
	FONT_TABLE_STYLE.name = 'Times New Roman'
	FONT_TABLE_STYLE.size = Pt(10)	

	print(type(document), document)

	for route in routes:

		document = get_docx(summary=route['summary'],
							settings=route['settings'],
							gen_time=datetime.strptime(route['start_time'], '%Y-%m-%d %H:%M:%S'),
							mass_export=True,
							document=document,
							TP_STYLE=TP_STYLE,
							HP_STYLE=HP_STYLE,
							TABLE_STYLE=TABLE_STYLE)

	name = 'gen_date' + '.docx'
	filename = os.path.join(app.config['UPLOAD_FOLDER'], name)
	document.save(filename)

	return name 